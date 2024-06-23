from docx import Document
from docx.shared import Pt, RGBColor, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT

doc = Document('abc.docx')

#поля
lst = doc.sections[0]
lst.left_margin = Mm(25)
lst.right_margin = Mm(10)
lst.top_margin = Mm(20)
lst.bottom_margin = Mm(20)

#списки НЕ ОТРАБАТЫВВЕТ НА МАРКИРОВАННЫХ
alf = tuple('абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ')
znk = tuple(',.')
for cpck in doc.paragraphs:
    if cpck.style.name.startswith('List') or cpck.style.name.startswith('List Bullet'):
        if cpck.text.endswith((alf)):
            cpck.text += ";"
        if cpck.text.endswith((znk)):
            cpck.text = cpck.text[:-1]+';'

#подписывание таблиц
coltblc = 1
for tblc in doc.tables:
    nwpar = doc.add_paragraph()
    nwpar.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    txtblc = nwpar.add_run(f'Таблица {coltblc} – ')
    ttbl = tblc._element
    ttbl.addprevious(nwpar._element)
    coltblc += 1


#текст
for txt in doc.paragraphs:
    for run in txt.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        txt.paragraph_format.line_spacing = 1.5
        txt.paragraph_format.space_before = Mm(0)
        txt.paragraph_format.space_after = Mm(0)
        txt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        txt.paragraph_format.first_line_indent = Cm(1.5)

#таблицы
for tblc in doc.tables:
    for row in tblc.rows:
        prw = tblc.rows[0]
        for cell in prw.cells:
            for txttblc in cell.paragraphs:
                for run in txttblc.runs:
                    run.font.bold = True
                    txttblc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            for txttblc in cell.paragraphs:
                for run in txttblc.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    txttblc.paragraph_format.line_spacing = 1
                    txttblc.paragraph_format.space_before = Mm(0)
                    txttblc.paragraph_format.space_after = Mm(0)
                    txttblc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#рисунки
colric = 1
for txt in doc.paragraphs:
    for run in txt.runs:
        if 'graphic' in run._element.xml:
            nwpar = doc.add_paragraph()
            nwpar.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            tric = nwpar.add_run(f'Рисунок {colric} – ')
            tric.font.size = Pt(14)
            tric.font.name = 'Times New Roman'
            rcnk = txt._element
            rcnk.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            txt.paragraph_format.first_line_indent = Cm(0)
            rcnk.addnext(nwpar._element)
            colric += 1


doc.save('test1.docx')